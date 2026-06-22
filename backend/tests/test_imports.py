"""Tests for imports endpoints: file upload, size limits, AI auto import."""
import io
from unittest.mock import patch

import pytest


def _make_docx_bytes(text: str = "Test content paragraph.\nSecond paragraph.") -> bytes:
    """Create a minimal .docx file in-memory."""
    from docx import Document
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestFileImport:
    FILE_UPLOAD = "/imports/file"
    FILE_AUTO = "/imports/file/auto"
    QUESTIONS_BATCH = "/questions/batch"

    def test_upload_docx(self, client, auth_headers):
        content = _make_docx_bytes("这是测试文本。\n第二行内容。")
        resp = client.post(
            self.FILE_UPLOAD,
            headers=auth_headers,
            files={"file": ("test.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "text" in data
        assert "这是测试文本" in data["text"]

    def test_upload_docx_returns_suggested_course_name(self, client, auth_headers):
        """File upload should return suggested_course_name derived from the filename."""
        content = _make_docx_bytes("Some content.")
        resp = client.post(
            self.FILE_UPLOAD,
            headers=auth_headers,
            files={"file": ("java复习题.docx", content,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["suggested_course_name"] == "java复习题"

    def test_upload_unsupported_format(self, client, auth_headers):
        resp = client.post(
            self.FILE_UPLOAD,
            headers=auth_headers,
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert resp.status_code == 400
        assert "不支持" in resp.json()["detail"]

    def test_upload_file_too_large(self, client, auth_headers):
        # Create content slightly larger than 10MB limit
        large_content = b"x" * (11 * 1024 * 1024)
        resp = client.post(
            self.FILE_UPLOAD,
            headers=auth_headers,
            files={"file": ("large.docx", large_content,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 413
        assert "超过" in resp.json()["detail"]

    def test_unauthorized(self, client):
        resp = client.post(self.FILE_UPLOAD, files={"file": ("test.docx", b"x", "application/octet-stream")})
        assert resp.status_code == 401

    @patch("backend.routers.imports.OPENAI_API_KEY", "")
    def test_auto_no_openai_key(self, client, auth_headers):
        """Without OPENAI_API_KEY set, /file/auto should return 400 with clear message."""
        content = _make_docx_bytes("Test question content.")
        resp = client.post(
            self.FILE_AUTO,
            headers=auth_headers,
            files={"file": ("test.docx", content,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 400
        detail = resp.json()["detail"]
        assert "OPENAI_API_KEY" in detail


class TestDeriveCourseName:
    """Tests for the derive_course_name_from_filename utility."""

    def test_normal_file(self):
        from backend.crud import derive_course_name_from_filename
        assert derive_course_name_from_filename("java复习题.docx") == "java复习题"

    def test_pptx_file(self):
        from backend.crud import derive_course_name_from_filename
        assert derive_course_name_from_filename("期末考试题.pptx") == "期末考试题"

    def test_whitespace_collapsed(self):
        from backend.crud import derive_course_name_from_filename
        assert derive_course_name_from_filename("  期末  考试 题.pptx ") == "期末 考试 题"

    def test_empty_filename(self):
        from backend.crud import derive_course_name_from_filename
        assert derive_course_name_from_filename("") == "未分类题库"

    def test_none_filename(self):
        from backend.crud import derive_course_name_from_filename
        assert derive_course_name_from_filename(None) == "未分类题库"

    def test_illegal_chars_stripped(self):
        from backend.crud import derive_course_name_from_filename
        result = derive_course_name_from_filename("math: test? <foo>.docx")
        assert "math test foo" == result or "math  test  foo" == result

    def test_truncate_long(self):
        from backend.crud import derive_course_name_from_filename
        long_name = "a" * 200 + ".docx"
        result = derive_course_name_from_filename(long_name)
        assert len(result) == 80


class TestCourseNameInBatchImport:
    """Test that /questions/batch respects course_name query param."""

    def test_batch_with_course_name_creates_new(self, client, auth_headers):
        """batch?course_name=Java复习题 should create that bank."""
        resp = client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Q?", "answer": "A",
        }], headers=auth_headers, params={"course_name": "Java复习题"})
        assert resp.status_code == 200
        data = resp.json()
        assert data["course_name"] == "Java复习题"
        assert data["course_id"] is not None

    def test_batch_with_course_name_reuses_existing(self, client, auth_headers):
        """Second batch with same course_name should reuse the bank."""
        client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Q1?", "answer": "A",
        }], headers=auth_headers, params={"course_name": "Reused"})
        resp2 = client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Q2?", "answer": "B",
        }], headers=auth_headers, params={"course_name": "Reused"})
        assert resp2.status_code == 200
        data = resp2.json()
        assert data["course_name"] == "Reused"

        # Should only be one "Reused" bank
        courses = client.get("/courses/mine", headers=auth_headers).json()
        matching = [c for c in courses if c["name"] == "Reused"]
        assert len(matching) == 1

    def test_batch_course_id_overrides_course_name(self, client, auth_headers):
        """course_id > 0 should take priority over course_name."""
        # Create a named course
        bank = client.post("/courses/", json={"name": "ExistingCourse"}, headers=auth_headers).json()
        cid = bank["id"]

        # Use course_id, passing a different course_name
        resp = client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Q?", "answer": "A",
        }], headers=auth_headers, params={"course_id": cid, "course_name": "Ignored"})
        assert resp.status_code == 200
        assert resp.json()["course_name"] == "ExistingCourse"

    def test_batch_without_course_name_fallsback_uncategorized(self, client, auth_headers):
        """No course_id and no course_name → fallback to 未分类题库."""
        resp = client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Q?", "answer": "A",
        }], headers=auth_headers)
        assert resp.status_code == 200
        data = resp.json()
        assert data["course_name"] == "未分类题库"


class TestFileAutoCourseName:
    """Tests for /imports/file/auto course_name parameter."""

    FILE_AUTO = "/imports/file/auto"

    def _mock_openai_response(self):
        """Return a mock OpenAI response that returns a valid questions JSON."""
        import json
        mock = patch("backend.routers.imports.OpenAI")
        mock_client = mock.start()
        instance = mock_client.return_value
        instance.chat.completions.create.return_value.choices = [
            type("obj", (), {"message": type("obj", (), {"content": json.dumps({
                "questions": [{"type": "fill_blank", "question": "Mock Q?", "answer": "A"}]
            })})()})()
        ]
        return mock

    def _make_docx(self):
        return _make_docx_bytes("Some content.")

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_auto_course_name_creates_new(self, client, auth_headers):
        """course_name=Java复习题 should create that bank."""
        mock = self._mock_openai_response()
        try:
            content = self._make_docx()
            resp = client.post(
                self.FILE_AUTO,
                headers=auth_headers,
                files={"file": ("ignored.docx", content,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                params={"course_name": "Java复习题"},
            )
            assert resp.status_code == 200
            data = resp.json()
            assert data["course_name"] == "Java复习题"
            assert data["course_id"] is not None
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_auto_course_id_overrides_course_name(self, client, auth_headers):
        """course_id > 0 should take priority even when course_name is given."""
        # Create a course first
        bank = client.post("/courses/", json={"name": "ExistingCourse"}, headers=auth_headers).json()
        cid = bank["id"]

        mock = self._mock_openai_response()
        try:
            content = self._make_docx()
            resp = client.post(
                self.FILE_AUTO,
                headers=auth_headers,
                files={"file": ("ignored.docx", content,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                params={"course_id": cid, "course_name": "IgnoredName"},
            )
            assert resp.status_code == 200
            data = resp.json()
            assert data["course_name"] == "ExistingCourse"
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_auto_derives_from_filename_when_no_course_name(self, client, auth_headers):
        """With no course_name, should derive from filename."""
        mock = self._mock_openai_response()
        try:
            content = self._make_docx()
            resp = client.post(
                self.FILE_AUTO,
                headers=auth_headers,
                files={"file": ("java复习题.docx", content,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert resp.status_code == 200
            data = resp.json()
            assert data["course_name"] == "java复习题"
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_auto_user_isolation_same_name(self, client, auth_headers):
        """Different users with same course_name should get separate banks."""
        alice = auth_headers
        # Register bob
        client.post("/auth/register", json={
            "username": "bob_auto_cn", "password": "pass", "invite_code": "dev-invite",
        })
        r = client.post("/auth/login", json={"username": "bob_auto_cn", "password": "pass"})
        bob = {"Authorization": f"Bearer {r.json()['access_token']}"}

        mock = self._mock_openai_response()
        try:
            content = self._make_docx()

            r1 = client.post(
                self.FILE_AUTO, headers=alice,
                files={"file": ("ignored.docx", content, "application/octet-stream")},
                params={"course_name": "SharedName"},
            )
            r2 = client.post(
                self.FILE_AUTO, headers=bob,
                files={"file": ("ignored.docx", content, "application/octet-stream")},
                params={"course_name": "SharedName"},
            )
            assert r1.status_code == 200
            assert r2.status_code == 200
            assert r1.json()["course_id"] != r2.json()["course_id"]
        finally:
            mock.stop()


class TestCourseNameUserIsolation:
    """Different users should have separate banks even with the same name."""

    def _auth(self, client, username):
        client.post("/auth/register", json={
            "username": username, "password": "pass", "invite_code": "dev-invite",
        })
        r = client.post("/auth/login", json={"username": username, "password": "pass"})
        return {"Authorization": f"Bearer {r.json()['access_token']}"}

    def test_same_name_different_users(self, client):
        alice = self._auth(client, "alice_cn")
        bob = self._auth(client, "bob_cn")

        # Alice creates a course via batch import
        r = client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Alice Q?", "answer": "A",
        }], headers=alice, params={"course_name": "MyCourse"})
        alice_cid = r.json()["course_id"]

        # Bob does the same — should create a separate bank
        r = client.post("/questions/batch", json=[{
            "type": "fill_blank", "question": "Bob Q?", "answer": "B",
        }], headers=bob, params={"course_name": "MyCourse"})
        bob_cid = r.json()["course_id"]

        # Different course IDs
        assert alice_cid != bob_cid

        # Alice sees 1 course, Bob sees 1 course (their own)
        assert len(client.get("/courses/mine", headers=alice).json()) == 1
        assert len(client.get("/courses/mine", headers=bob).json()) == 1


class TestPreviewImport:
    """Tests for /imports/file/preview — parse only, no DB writes."""

    PREVIEW_URL = "/imports/file/preview"

    def _mock_openai_response(self, questions_data=None):
        """Return a mock OpenAI response for preview."""
        import json
        if questions_data is None:
            questions_data = [{"type": "fill_blank", "question": "Preview Q?", "answer": "Ans"}]
        mock = patch("backend.routers.imports.OpenAI")
        mc = mock.start()
        inst = mc.return_value
        inst.chat.completions.create.return_value.choices = [
            type("obj", (), {"message": type("obj", (), {"content": json.dumps({
                "questions": questions_data
            })})()})()
        ]
        return mock

    def _make_docx(self, text="Content."):
        return _make_docx_bytes(text)

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_preview_returns_questions_no_db_write(self, client, auth_headers):
        """Preview should return parsed questions and NOT create any DB records."""
        mock = self._mock_openai_response()
        try:
            resp = client.post(
                self.PREVIEW_URL, headers=auth_headers,
                files={"file": ("test.docx", self._make_docx(), "application/octet-stream")},
            )
            assert resp.status_code == 200
            data = resp.json()
            assert len(data["questions"]) == 1
            assert data["total_parsed"] == 1
            assert data["total_valid"] == 1
            assert data["suggested_course_name"] == "test"
            assert data["timing"]["extract_ms"] >= 0
            assert data["timing"]["ai_ms"] >= 0
            assert data["timing"]["total_ms"] >= 0
            assert data["timing"]["chunks"] == 1
            # Verify no questions were created in DB
            q_resp = client.get("/questions/", headers=auth_headers)
            assert len(q_resp.json()) == 0
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_preview_returns_validation_warnings(self, client, auth_headers):
        """Preview should flag invalid questions without crashing."""
        data = [
            {"type": "invalid_type", "question": "Bad?", "answer": "X"},
            {"type": "fill_blank", "question": "Good?", "answer": "Y"},
        ]
        mock = self._mock_openai_response(data)
        try:
            resp = client.post(
                self.PREVIEW_URL, headers=auth_headers,
                files={"file": ("test.docx", self._make_docx(), "application/octet-stream")},
            )
            assert resp.status_code == 200
            d = resp.json()
            # Only valid question returned
            assert len(d["questions"]) == 1
            assert d["questions"][0]["question"] == "Good?"
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_preview_no_questions_returns_clear_error(self, client, auth_headers):
        """When AI returns no questions, preview should fail with a clear message."""
        mock = self._mock_openai_response([])
        try:
            resp = client.post(
                self.PREVIEW_URL, headers=auth_headers,
                files={"file": ("test.docx", self._make_docx(), "application/octet-stream")},
            )
            assert resp.status_code == 400
            assert "未能从文档中解析出任何有效题目" in resp.json()["detail"]
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "")
    def test_preview_no_openai_key(self, client, auth_headers):
        """Without key, preview returns 400."""
        resp = client.post(
            self.PREVIEW_URL, headers=auth_headers,
            files={"file": ("test.docx", self._make_docx(), "application/octet-stream")},
        )
        assert resp.status_code == 400
        assert "OPENAI_API_KEY" in resp.json()["detail"]

    def test_preview_unsupported_format(self, client, auth_headers):
        resp = client.post(
            self.PREVIEW_URL, headers=auth_headers,
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert resp.status_code == 400
        assert "不支持" in resp.json()["detail"]


class TestConfirmImport:
    """Tests for /imports/confirm — single-transaction import."""

    CONFIRM_URL = "/imports/confirm"

    def test_confirm_creates_questions(self, client, auth_headers):
        """Confirm with valid questions should create questions and a course."""
        resp = client.post(self.CONFIRM_URL, json={
            "questions": [
                {"type": "fill_blank", "question": "Confirm Q?", "answer": "Ans"},
                {"type": "single_choice", "question": "Choice?",
                 "options": {"A": "a", "B": "b"}, "answer": "A"},
            ],
            "course_name": "ConfirmCourse",
        }, headers=auth_headers)
        assert resp.status_code == 200
        data = resp.json()
        assert data["imported_count"] == 2
        assert data["course_name"] == "ConfirmCourse"
        assert data["course_id"] is not None

        # Verify DB has them
        qs = client.get("/questions/", headers=auth_headers).json()
        assert len(qs) == 2

    def test_confirm_existing_course(self, client, auth_headers):
        """Confirm with existing course_id should add to that course."""
        bank = client.post("/courses/", json={"name": "Existing"}, headers=auth_headers).json()
        cid = bank["id"]
        resp = client.post(self.CONFIRM_URL, json={
            "questions": [
                {"type": "fill_blank", "question": "Add to existing?", "answer": "Yes"},
            ],
            "course_id": cid,
        }, headers=auth_headers)
        assert resp.status_code == 200
        assert resp.json()["imported_count"] == 1
        assert resp.json()["course_name"] == "Existing"

    def test_confirm_rollback_on_invalid(self, client, auth_headers):
        """If one question is invalid, no course or questions should be created."""
        resp = client.post(self.CONFIRM_URL, json={
            "questions": [
                {"type": "fill_blank", "question": "Valid Q?", "answer": "Correct"},
                {"type": "single_choice", "question": "No options?",
                 "options": None, "answer": "A"},
            ],
            "course_name": "RollbackTest",
        }, headers=auth_headers)
        assert resp.status_code == 422

        # Verify no questions were created
        qs = client.get("/questions/", headers=auth_headers).json()
        assert len(qs) == 0

        # Verify no course was created (validation happens before course resolution)
        courses = client.get("/courses/mine", headers=auth_headers).json()
        assert all(c["name"] != "RollbackTest" for c in courses)

    def test_confirm_empty_questions_rejected(self, client, auth_headers):
        resp = client.post(self.CONFIRM_URL, json={
            "questions": [],
            "course_name": "Empty",
        }, headers=auth_headers)
        assert resp.status_code == 422
        assert "没有" in resp.text

    def test_confirm_nonexistent_course_id(self, client, auth_headers):
        resp = client.post(self.CONFIRM_URL, json={
            "questions": [
                {"type": "fill_blank", "question": "Q?", "answer": "A"},
            ],
            "course_id": 99999,
        }, headers=auth_headers)
        assert resp.status_code == 400


class TestEnhancedTextExtraction:
    """Tests for enhanced text extraction: tables, images."""

    def test_validator_rejects_bad_question(self):
        """Test _validate_question_item directly."""
        from backend.routers.imports import _validate_question_item
        # Missing question
        result, err = _validate_question_item({"type": "fill_blank", "answer": "A"})
        assert result is None
        assert err is not None

    def test_validator_accepts_good_question(self):
        from backend.routers.imports import _validate_question_item
        result, err = _validate_question_item({
            "type": "single_choice", "question": "Q?",
            "options": {"A": "a", "B": "b"}, "answer": "A"
        })
        assert result is not None
        assert err is None
        assert result["type"] == "single_choice"

    def test_validator_rejects_choice_without_options(self):
        from backend.routers.imports import _validate_question_item
        result, err = _validate_question_item({
            "type": "single_choice", "question": "Q?", "answer": "A"
        })
        assert result is None
        assert "选项" in (err or "")

    def test_docx_table_text_extraction(self, tmp_path):
        """Verify docx tables are included in extracted text."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Normal paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Q1"
        table.cell(0, 1).text = "Answer1"
        table.cell(1, 0).text = "Q2"
        table.cell(1, 1).text = "Answer2"
        path = str(tmp_path / "table_test.docx")
        doc.save(path)
        from backend.routers.imports import extract_text_and_warnings
        text, warnings = extract_text_and_warnings(path)
        assert "Normal paragraph" in text
        assert "Q1 | Answer1" in text
        assert "Q2 | Answer2" in text


class TestAutoImportAIFailure:
    """AI failure in /imports/file/auto: course may be created, but no questions added."""

    FILE_AUTO = "/imports/file/auto"

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_ai_empty_questions_returns_400_no_questions_created(self, client, auth_headers):
        """When AI returns zero questions, endpoint returns 400 and no questions in DB."""
        import json
        mock = patch("backend.routers.imports.OpenAI")
        mc = mock.start()
        try:
            inst = mc.return_value
            inst.chat.completions.create.return_value.choices = [
                type("obj", (), {"message": type("obj", (), {"content": json.dumps({
                    "questions": []
                })})()})()
            ]
            content = _make_docx_bytes("Some content that yields no questions.")
            resp = client.post(
                self.FILE_AUTO, headers=auth_headers,
                files={"file": ("test.docx", content, "application/octet-stream")},
            )
            # Endpoint returns 400 — no questions parsed
            assert resp.status_code == 400
            assert "解析" in resp.json()["detail"]
            # No questions were created in DB
            qs = client.get("/questions/", headers=auth_headers).json()
            assert len(qs) == 0
        finally:
            mock.stop()

    @patch("backend.routers.imports.OPENAI_API_KEY", "sk-test")
    def test_ai_bad_json_returns_400_no_questions_created(self, client, auth_headers):
        """When AI returns malformed JSON, endpoint returns 400 and no questions."""
        mock = patch("backend.routers.imports.OpenAI")
        mc = mock.start()
        try:
            inst = mc.return_value
            inst.chat.completions.create.return_value.choices = [
                type("obj", (), {"message": type("obj", (), {"content": "this is not json"})()})()
            ]
            content = _make_docx_bytes("Content.")
            resp = client.post(
                self.FILE_AUTO, headers=auth_headers,
                files={"file": ("test.docx", content, "application/octet-stream")},
            )
            assert resp.status_code >= 400
            qs = client.get("/questions/", headers=auth_headers).json()
            assert len(qs) == 0
        finally:
            mock.stop()
